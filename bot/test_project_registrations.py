"""Test script cho project_registrations API"""
import asyncio
import json
import sys

# Test data - sample proposal files
SAMPLE_FILES_JSON = json.dumps([
    {"name": "proposal.docx", "url": "/uploads/proposal-1.docx"},
    {"name": "report.pdf", "url": "/uploads/report-1.pdf"}
])

async def test_extract_text_from_doc():
    """Test hàm _extract_text_from_doc"""
    from src.api.routes.project_registrations import _extract_text_from_doc
    
    # Test với bytes rỗng để xem có lỗi gì không
    print("\n=== Test _extract_text_from_doc ===")
    
    try:
        # Tạo một DOCX đơn giản để test
        from docx import Document
        from io import BytesIO
        
        # Tạo DOCX trong memory
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test paragraph.')
        doc.add_paragraph('Another paragraph with more text.')
        
        # Save vào bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_bytes = buffer.read()
        
        print(f"Created test DOCX: {len(docx_bytes)} bytes")
        
        # Test extract
        result = await _extract_text_from_doc(docx_bytes, "test.docx")
        print(f"Result: {result}")
        
        if result:
            print("✓ DOCX extraction successful!")
        else:
            print("✗ DOCX extraction returned None")
            
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback
        traceback.print_exc()


async def test_ocr_file_from_url():
    """Test hàm _ocr_file_from_url"""
    from src.api.routes.project_registrations import _ocr_file_from_url
    
    print("\n=== Test _ocr_file_from_url ===")
    
    # Test với URL không tồn tại - xem có log ra lỗi gì
    try:
        result = await _ocr_file_from_url("/nonexistent/file.docx", "test.docx")
        print(f"Result for non-existent file: {result}")
    except Exception as e:
        print(f"Error (expected for non-existent URL): {e}")


async def test_extract_ocr_text():
    """Test hàm _extract_ocr_text"""
    from src.api.routes.project_registrations import _extract_ocr_text
    
    print("\n=== Test _extract_ocr_text ===")
    
    result = await _extract_ocr_text(SAMPLE_FILES_JSON)
    print(f"Result: {result}")


async def main():
    print("Testing project_registrations API functions...")
    
    await test_extract_text_from_doc()
    await test_ocr_file_from_url()
    await test_extract_ocr_text()
    
    print("\n=== Done ===")


if __name__ == "__main__":
    asyncio.run(main())
