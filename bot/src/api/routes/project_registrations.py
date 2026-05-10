"""
ProjectRegistration API Routes - Xử lý lấy thông tin đăng ký đề tài
"""

import json
import os
from datetime import datetime
from typing import Any

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel, Field

from src.clients.embedding_client import embed_text
from src.clients.qdrant_client import check_collection_exists, create_collection, upsert_points
from src.db import db
from src.repositories.project_registration import get_project_registration_by_call_round_id
from src.services.llm_service import LLMService
from src.utilities import get_logger, log_async_execution

logger = get_logger(__name__)
router = APIRouter(prefix="/api/project-registrations", tags=["project-registrations"])


class GetByIdsRequest(BaseModel):
    """Request model cho việc lấy danh sách đăng ký theo IDs"""
    ids: list[str] = Field(..., description="Danh sách ID của các bản ghi đăng ký")


class CallRoundFilters(BaseModel):
    callRoundId: str = Field(..., description="ID của đợt đăng ký")
    fromDate: str | None = Field(None, description="ISO datetime từ ngày")
    toDate: str | None = Field(None, description="ISO datetime đến ngày")


class EvaluationCriteria(BaseModel):
    minScore: int = Field(70, description="Điểm tối thiểu để APPROVE")
    requireInstructor: bool = Field(True, description="Bắt buộc có giảng viên hướng dẫn")
    checkPlagiarism: bool = Field(False, description="Có kiểm tra đạo văn hay không")


class EvaluateByCallRoundRequest(BaseModel):
    filters: CallRoundFilters
    criteria: EvaluationCriteria


class IndexApprovedRequest(BaseModel):
    callRoundId: str = Field(..., description="ID của đợt đăng ký")
    approvedIds: list[str] = Field(..., description="Danh sách ProjectRegistration IDs được duyệt")


@router.post("/batch")
@log_async_execution
async def get_project_registrations_by_ids(request: GetByIdsRequest) -> dict[str, Any]:
    """
    Lấy thông tin các bản ghi đăng ký đề tài dựa trên danh sách ID
    
    Args:
        request: Object chứa danh sách IDs cần lấy
        
    Returns:
        Dict chứa danh sách bản ghi đăng ký
    """
    logger.info(f"Lấy thông tin đăng ký đề tài - Số lượng IDs: {len(request.ids)}")
    
    if not request.ids:
        logger.warning("Danh sách IDs rỗng")
        return {
            "success": True,
            "data": [],
            "total": 0,
            "message": "Danh sách IDs rỗng"
        }
    
    try:
        # Chuyển đổi list IDs thành placeholder $1, $2, ...
        placeholders = ", ".join([f"${i+1}" for i in range(len(request.ids))])
        
        query = f'''
            SELECT 
                pr.id,
                pr."userId",
                pr."callRoundId",
                pr.title,
                pr.objective,
                pr."expectedOutput",
                pr."proposalFiles",
                pr."teamMembers",
                pr.status,
                pr."cancelReason",
                pr."createdAt",
                pr."updatedAt",
                pr."instructorId",
                pr."instructorStatus",
                pr."facultyStatus",
                pr."facultyReviewerId",
                u.name as "userName",
                u.email as "userEmail",
                cr.name as "callRoundName"
            FROM "ProjectRegistration" pr
            LEFT JOIN "User" u ON pr."userId" = u.id
            LEFT JOIN "CallRound" cr ON pr."callRoundId" = cr.id
            WHERE pr.id IN ({placeholders})
            ORDER BY pr."createdAt" DESC
        '''
        
        rows = await db.fetch(query, *request.ids)
        
        # Chuyển đổi kết quả thành list dict
        registrations = []
        for row in rows:
            reg = dict(row)
            # Chuyển đổi datetime objects thành ISO string
            if reg.get("createdAt"):
                reg["createdAt"] = reg["createdAt"].isoformat()
            if reg.get("updatedAt"):
                reg["updatedAt"] = reg["updatedAt"].isoformat()
            
            # Trích xuất fullText từ OCR file đính kèm
            full_text = await _extract_ocr_text(reg.get("proposalFiles"))
            reg["fullText"] = full_text
            
            registrations.append(reg)
        
        # Index documents to Qdrant - collection name based on CallRound.name
        from src.clients.embedding_client import embed_text
        from src.clients.qdrant_client import check_collection_exists, create_collection, upsert_point
        
        # sentence-transformers/all-MiniLM-L6-v2 produces 384-dim vectors
        VECTOR_SIZE = int(os.getenv("EMBEDDING_VECTOR_SIZE", "384"))
        
        # Lấy CallRound names để tạo collection names
        call_round_ids = list(set(r.get("callRoundId") for r in registrations if r.get("callRoundId")))
        call_round_names = {}
        
        if call_round_ids:
            placeholders = ", ".join([f"${i+1}" for i in range(len(call_round_ids))])
            cr_query = f'SELECT id, name FROM "CallRound" WHERE id IN ({placeholders})'
            cr_rows = await db.fetch(cr_query, *call_round_ids)
            call_round_names = {row["id"]: row["name"] for row in cr_rows}
            logger.info(f"Tìm thấy {len(call_round_names)} CallRound: {call_round_names}")
        
        def _sanitize_collection_name(name: str) -> str:
            """Sanitize CallRound name thành valid Qdrant collection name"""
            sanitized = name.strip().replace(" ", "_").replace("-", "_")
            sanitized = ''.join(c for c in sanitized if c.isalnum() or c == "_")
            return sanitized[:50] if sanitized else "default"
        
        for reg in registrations:
            try:
                reg_id = reg.get("id")
                if not reg_id:
                    continue
                
                # Build collection name từ CallRound.name
                call_round_id = reg.get("callRoundId")
                call_round_name = call_round_names.get(call_round_id) if call_round_id else None
                collection_name = _sanitize_collection_name(call_round_name) if call_round_name else "default"
                
                # Ensure collection exists
                if not check_collection_exists(collection_name):
                    create_collection(collection_name, VECTOR_SIZE)
                    logger.info(f"Tạo collection mới: {collection_name}")
                
                # Build fullText từ title + objective + ocr text
                full_text = reg.get("fullText") or ""
                title = reg.get("title") or ""
                objective = reg.get("objective") or ""
                
                combined_text = f"{title}\n{objective}\n{full_text}".strip()
                
                if combined_text:
                    vector = await embed_text(combined_text)
                    if vector:
                        payload = {
                            "id": reg_id,
                            "projectId": reg_id,
                            "title": title,
                            "objective": objective,
                            "status": reg.get("status"),
                            "userId": reg.get("userId"),
                            "callRoundId": call_round_id,
                            "callRoundName": call_round_name,
                            "userName": reg.get("userName"),
                            "fullText": reg.get("fullText")  # Limit size
                        }
                        if upsert_point(collection_name, reg_id, vector, payload):
                            logger.info(f"Indexed project {reg_id} vào collection '{collection_name}'")
                        else:
                            logger.warning(f"Failed to index project {reg_id} into collection '{collection_name}'")
            except Exception as exc:
                logger.warning(f"Failed to index project {reg.get('id')}: {exc}")
        
        logger.info(f"Đã lấy {len(registrations)} bản ghi đăng ký đề tài")
        
        return {
            "success": True,
            "data": registrations,
            "total": len(registrations),
            "message": f"Tìm thấy {len(registrations)} bản ghi"
        }
        
    except Exception as exc:
        logger.error(f"Lỗi khi lấy thông tin đăng ký: {str(exc)}")
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/evaluate-by-call-round")
@log_async_execution
async def evaluate_by_call_round(request: EvaluateByCallRoundRequest) -> dict[str, Any]:
    """
    Đánh giá ProjectRegistration theo callRoundId, OCR file PDF và dùng LLM.
    """
    call_round_id = request.filters.callRoundId
    criteria = request.criteria.model_dump()

    try:
        registrations = await get_project_registration_by_call_round_id(call_round_id)

        from_dt = _parse_datetime(request.filters.fromDate)
        to_dt = _parse_datetime(request.filters.toDate)

        if from_dt or to_dt:
            registrations = [
                r for r in registrations
                if _is_in_date_range(r.get("createdAt"), from_dt, to_dt)
            ]

        if not registrations:
            return {
                "summary": {
                    "total": 0,
                    "approved": 0,
                    "revision": 0,
                    "rejected": 0,
                    "errors": 0,
                },
                "evaluations": [],
            }

        llm_service = LLMService()
        evaluations: list[dict[str, Any]] = []

        for reg in registrations:
            reg["ocrFullText"] = await _extract_ocr_text(reg.get("proposalFiles"))
            ocr_matches = _compare_fields_to_ocr(reg, reg.get("ocrFullText"))

            evaluation = await _evaluate_project_with_ocr(
                reg,
                criteria,
                llm_service,
                ocr_matches,
            )

            evaluations.append(evaluation)

        await _apply_duplicate_rules(registrations, evaluations)

        summary = {
            "total": len(evaluations),
            "approved": sum(1 for e in evaluations if e.get("decision") == "APPROVE"),
            "revision": sum(1 for e in evaluations if e.get("decision") == "REVISION"),
            "rejected": sum(1 for e in evaluations if e.get("decision") == "REJECT"),
            "errors": sum(1 for e in evaluations if e.get("decision") == "ERROR"),
        }

        return {
            "summary": summary,
            "evaluations": evaluations,
        }

    except Exception as exc:
        logger.error(f"Lỗi khi đánh giá project theo call round: {exc}")
        raise HTTPException(status_code=500, detail=str(exc))


@router.post("/index-approved")
@log_async_execution
async def index_approved_projects(request: IndexApprovedRequest) -> dict[str, Any]:
    """
    Index các bản ghi được duyệt vào Qdrant theo collection của CallRound.
    """
    if not request.approvedIds:
        return {
            "success": True,
            "indexed": 0,
            "message": "Danh sách approvedIds rỗng",
        }

    try:
        call_round_name = await _get_call_round_name(request.callRoundId)
        collection_name = _sanitize_collection_name(call_round_name) if call_round_name else "default"

        vector_size = int(os.getenv("EMBEDDING_VECTOR_SIZE", "384"))
        if not check_collection_exists(collection_name):
            create_collection(collection_name, vector_size)

        placeholders = ", ".join([f"${i+1}" for i in range(len(request.approvedIds))])
        query = f'''
            SELECT 
                pr.id,
                pr."userId",
                pr."callRoundId",
                pr.title,
                pr.objective,
                pr."expectedOutput",
                pr."proposalFiles",
                pr."teamMembers",
                pr.status,
                pr."createdAt",
                pr."updatedAt",
                pr."instructorId",
                u.name as "userName",
                u.email as "userEmail",
                cr.name as "callRoundName"
            FROM "ProjectRegistration" pr
            LEFT JOIN "User" u ON pr."userId" = u.id
            LEFT JOIN "CallRound" cr ON pr."callRoundId" = cr.id
            WHERE pr.id IN ({placeholders})
            ORDER BY pr."createdAt" DESC
        '''

        rows = await db.fetch(query, *request.approvedIds)

        points: list[dict[str, Any]] = []
        for row in rows:
            reg = dict(row)
            ocr_text = await _extract_ocr_text(reg.get("proposalFiles"))
            title = reg.get("title") or ""
            objective = reg.get("objective") or ""
            combined_text = f"{title}\n{objective}\n{ocr_text or ''}".strip()
            if not combined_text:
                continue

            vector = await embed_text(combined_text)
            if not vector:
                continue

            points.append({
                "id": reg.get("id"),
                "vector": vector,
                "payload": {
                    "id": reg.get("id"),
                    "projectId": reg.get("id"),
                    "title": title,
                    "objective": objective,
                    "status": reg.get("status"),
                    "userId": reg.get("userId"),
                    "callRoundId": reg.get("callRoundId"),
                    "callRoundName": reg.get("callRoundName"),
                    "userName": reg.get("userName"),
                    "userEmail": reg.get("userEmail"),
                },
            })

        if points:
            upsert_points(collection_name, points)

        return {
            "success": True,
            "indexed": len(points),
            "collection": collection_name,
        }

    except Exception as exc:
        logger.error(f"Lỗi khi index approved projects: {exc}")
        raise HTTPException(status_code=500, detail=str(exc))


def _build_fulltext_from_fields(title: str | None, objective: str | None) -> str | None:
    """Tạo fullText từ title và objective của đăng ký đề tài"""
    parts = []
    if title:
        parts.append(f"TIÊU ĐỀ: {title}")
    if objective:
        parts.append(f"MỤC TIÊU NGHIÊN CỨU: {objective}")
    if parts:
        return "\n\n".join(parts)
    return None


async def _extract_ocr_text(proposal_files_json: str | None) -> str | None:
    """Trích xuất text từ OCR của các file đính kèm trong đăng ký đề tài"""
    if not proposal_files_json:
        return None
    
    try:
        files = json.loads(proposal_files_json)
        if not files or not isinstance(files, list):
            return None
        
        full_text_parts = []
        for file_info in files:
            if not isinstance(file_info, dict):
                continue
            
            file_url = file_info.get("url")
            file_name = file_info.get("name", "unknown")
            
            if not file_url:
                continue
            
            ocr_text = await _ocr_file_from_url(file_url, file_name)
            if ocr_text:
                full_text_parts.append(ocr_text)
        
        if full_text_parts:
            return "\n\n---\n\n".join(full_text_parts)
        return None
        
    except json.JSONDecodeError as exc:
        logger.warning(f"Không parse được proposalFiles JSON: {exc}")
        return None
    except Exception as exc:
        logger.error(f"Lỗi khi trích xuất OCR text: {exc}")
        return None


async def _ocr_file_from_url(file_url: str, file_name: str) -> str | None:
    """Tải file từ URL và chạy OCR"""
    import httpx
    from src.services.ocr.ocr_service import run_ocr_with_vllm
    
    try:
        app_url = os.getenv("APP_URL")
        base_url = app_url if app_url else "http://localhost:3000"
        full_url = f"{base_url}{file_url}" if file_url.startswith("/") else file_url
        
        logger.info(f"Tải file: {full_url}")
        logger.info(f"  - file_name: {file_name}")
        logger.info(f"  - APP_URL: {app_url}")
        
        async with httpx.AsyncClient(timeout=60.0) as client:
            response = await client.get(full_url)
            response.raise_for_status()
            file_bytes = response.content
            logger.info(f"  - Tải thành công: {len(file_bytes)} bytes")
        
        lower_name = file_name.lower()
        
        # Chỉ xử lý file PDF
        if not lower_name.endswith(".pdf"):
            logger.debug(f"Bỏ qua file không phải PDF: {file_name}")
            return None
        
        logger.info(f"Chạy OCR cho file PDF: {file_name}")
        ocr_result = await run_ocr_with_vllm(file_bytes, file_name)
        
        if ocr_result and ocr_result.text:
            logger.info(f"  - OCR thành công: {len(ocr_result.text)} ký tự")
            return ocr_result.text
        
        logger.warning(f"  - OCR không trả về text")
        return None
        
    except httpx.HTTPStatusError as exc:
        logger.error(f"Không tải được file {file_url}: status={exc.response.status_code}")
        return None
    except Exception as exc:
        logger.error(f"Lỗi xử lý file {file_name}: {type(exc).__name__}: {exc}")
        return None


def _parse_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None


def _is_in_date_range(
    created_at: datetime | str | None,
    from_dt: datetime | None,
    to_dt: datetime | None,
) -> bool:
    if created_at is None:
        return True

    if isinstance(created_at, str):
        created_at = _parse_datetime(created_at)
    if created_at is None:
        return True

    if from_dt and created_at < from_dt:
        return False
    if to_dt and created_at > to_dt:
        return False
    return True


def _normalize_text(value: str | None) -> str:
    if not value:
        return ""
    return " ".join(value.lower().split())


def _compare_fields_to_ocr(reg: dict[str, Any], ocr_text: str | None) -> dict[str, bool]:
    normalized_ocr = _normalize_text(ocr_text)
    return {
        "title": _normalize_text(reg.get("title")) in normalized_ocr,
        "objective": _normalize_text(reg.get("objective")) in normalized_ocr,
        "expectedOutput": _normalize_text(reg.get("expectedOutput")) in normalized_ocr,
    }


def _decision_from_score(score: int, min_score: int) -> str:
    if score >= min_score:
        return "APPROVE"
    if score >= 50:
        return "REVISION"
    return "REJECT"


async def _evaluate_project_with_ocr(
    reg: dict[str, Any],
    criteria: dict[str, Any],
    llm_service: LLMService,
    ocr_matches: dict[str, bool],
) -> dict[str, Any]:
    mismatch_fields = [k for k, v in ocr_matches.items() if not v]
    min_score = int(criteria.get("minScore", 70))
    require_instructor = bool(criteria.get("requireInstructor", True))

    prompt = f"""Ban la chuyen gia danh gia de tai nghien cuu. Hay danh gia de tai dua tren thong tin nhap va text OCR.

Thong tin dang ky:
- Tieu de: {reg.get('title', 'N/A')}
- Muc tieu: {reg.get('objective', 'N/A')}
- Ket qua du kien: {reg.get('expectedOutput', 'N/A')}
- Giang vien huong dan: {reg.get('instructorName', 'N/A')}

Ket qua OCR (trich xuat tu file PDF):
{(reg.get('ocrFullText') or '')[:6000]}

Ket qua doi chieu OCR (true la co khop): {ocr_matches}
Tieu chi diem: min_score={min_score}

Yeu cau:
1. Danh gia diem 0-100
2. Quyết dinh: APPROVE/REVISION/REJECT
3. Neu OCR khong khop, can giam diem hoac yeu cau REVISION
4. Tra ve JSON: {{"score": <int>, "decision": "<APPROVE|REVISION|REJECT>", "reason": "<ly do>"}}
"""

    try:
        response = await llm_service.chat_completion(
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert evaluator. Always respond in Vietnamese and return valid JSON.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            response_format={"type": "json_object"},
        )

        result = json.loads(response)
        score = int(result.get("score", 0))
        decision = _decision_from_score(score, min_score)
        reason = result.get("reason", "")

        if mismatch_fields:
            reason = f"{reason} (OCR khong khop: {', '.join(mismatch_fields)})"

        if require_instructor and not reg.get("instructorId"):
            decision = "REJECT"
            reason = f"{reason} (Thieu giang vien huong dan)"

        return {
            "projectId": reg.get("id"),
            "registrationId": reg.get("id"),
            "projectTitle": reg.get("title"),
            "score": score,
            "decision": decision,
            "reason": reason,
            "ocrMatches": ocr_matches,
            "ocrFullText": reg.get("ocrFullText"),
            "isDuplicate": False,
            "duplicateScore": None,
            "duplicateOf": None,
            "evaluatedAt": datetime.utcnow().isoformat() + "Z",
        }

    except Exception as exc:
        logger.error(f"Lỗi khi đánh giá project {reg.get('id')}: {exc}")
        return {
            "projectId": reg.get("id"),
            "registrationId": reg.get("id"),
            "projectTitle": reg.get("title"),
            "score": 0,
            "decision": "ERROR",
            "reason": f"Loi khi danh gia: {str(exc)}",
            "ocrMatches": ocr_matches,
            "ocrFullText": reg.get("ocrFullText"),
            "isDuplicate": False,
            "duplicateScore": None,
            "duplicateOf": None,
            "evaluatedAt": datetime.utcnow().isoformat() + "Z",
        }


def _cosine_similarity(vec_a: list[float], vec_b: list[float]) -> float:
    if not vec_a or not vec_b or len(vec_a) != len(vec_b):
        return 0.0
    dot = sum(a * b for a, b in zip(vec_a, vec_b))
    norm_a = sum(a * a for a in vec_a) ** 0.5
    norm_b = sum(b * b for b in vec_b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


def _get_created_at(reg: dict[str, Any]) -> datetime | None:
    value = reg.get("createdAt")
    if isinstance(value, datetime):
        return value
    if isinstance(value, str):
        return _parse_datetime(value)
    return None


async def _apply_duplicate_rules(
    registrations: list[dict[str, Any]],
    evaluations: list[dict[str, Any]],
) -> None:
    threshold = float(os.getenv("QDRANT_DUPLICATE_THRESHOLD", "0.85"))
    reg_by_id = {r.get("id"): r for r in registrations}
    eval_by_id = {e.get("registrationId"): e for e in evaluations}

    embeddings: dict[str, list[float]] = {}
    for reg in registrations:
        reg_id = reg.get("id")
        text = reg.get("ocrFullText") or ""
        if not reg_id or not text.strip():
            continue
        try:
            embeddings[reg_id] = await embed_text(text)
        except Exception as exc:
            logger.warning(f"Khong tao duoc embedding cho {reg_id}: {exc}")

    ids = list(embeddings.keys())
    for i in range(len(ids)):
        for j in range(i + 1, len(ids)):
            left_id = ids[i]
            right_id = ids[j]
            left_vec = embeddings.get(left_id)
            right_vec = embeddings.get(right_id)
            if not left_vec or not right_vec:
                continue
            score = _cosine_similarity(left_vec, right_vec)
            if score < threshold:
                continue

            left_reg = reg_by_id.get(left_id)
            right_reg = reg_by_id.get(right_id)
            if not left_reg or not right_reg:
                continue

            left_created = _get_created_at(left_reg)
            right_created = _get_created_at(right_reg)
            keep_id = left_id
            dup_id = right_id
            if left_created and right_created and right_created < left_created:
                keep_id = right_id
                dup_id = left_id

            dup_eval = eval_by_id.get(dup_id)
            if dup_eval:
                dup_eval["isDuplicate"] = True
                dup_eval["duplicateScore"] = score
                dup_eval["duplicateOf"] = keep_id
                dup_eval["decision"] = "REJECT"
                dup_eval["reason"] = f"{dup_eval.get('reason', '')} (Trung lap voi {keep_id})"


async def _get_call_round_name(call_round_id: str) -> str | None:
    row = await db.fetchrow('SELECT name FROM "CallRound" WHERE id = $1', call_round_id)
    if not row:
        return None
    return row.get("name")


def _sanitize_collection_name(name: str) -> str:
    sanitized = name.strip().replace(" ", "_").replace("-", "_")
    sanitized = "".join(c for c in sanitized if c.isalnum() or c == "_")
    return sanitized[:50] if sanitized else "default"


async def _extract_text_from_doc(file_bytes: bytes, file_name: str) -> str | None:
    """Trích xuất text từ file DOC/DOCX"""
    import tempfile
    
    is_docx = file_name.lower().endswith('.docx')
    
    try:
        if is_docx:
            try:
                from docx import Document
                
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                
                try:
                    doc = Document(tmp_path)
                    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                    text = '\n'.join(paragraphs)
                    
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            if any(cells):
                                text += '\n' + ' | '.join(cells)
                    
                    logger.info(f"Đọc DOCX thành công - {len(text)} ký tự")
                    return text if text.strip() else None
                finally:
                    os.unlink(tmp_path)
                    
            except ImportError:
                logger.warning("python-docx chưa được cài đặt")
                return None
        else:
            try:
                import docx2txt
                
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                    tmp.write(file_bytes)
                    tmp_path = tmp.name
                
                try:
                    text = docx2txt.process(tmp_path)
                    logger.info(f"Đọc DOC bằng docx2txt thành công - {len(text)} ký tự")
                    return text if text.strip() else None
                finally:
                    os.unlink(tmp_path)
            except ImportError:
                logger.warning("docx2txt chưa được cài đặt")
                return None
            except Exception as exc:
                logger.warning(f"docx2txt lỗi: {exc}")
                return None
                
    except Exception as exc:
        logger.error(f"Lỗi trích xuất text từ DOC: {exc}")
        return None